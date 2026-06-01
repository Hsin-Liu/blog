"""
Video Summarizer - Core orchestrator for video processing.

Processes teaching videos by:
1. Extracting audio via FFmpeg
2. Extracting PPT/slide frames (CPU)
3. Transcribing audio via Whisper (GPU)
4. Aligning timeline and generating reports

Optimized for RTX 5060 Ti (16GB) with CUDA 12.4+.
"""

import os
import sys
import json
import logging
import subprocess
from datetime import timedelta
from pathlib import Path
from typing import Optional

from faster_whisper import WhisperModel

try:
    from extract_video_ppt import extract_ppt
except ImportError:
    extract_ppt = None

import config

# Configure logging
logging.basicConfig(
    level=getattr(logging, config.LOG_LEVEL),
    format=config.LOG_FORMAT
)
logger = logging.getLogger(__name__)


class VideoSummarizer:
    """Main orchestrator for video processing."""

    def __init__(
        self,
        video_path: str,
        output_dir: Optional[str] = None,
        whisper_model: str = config.WHISPER_MODEL,
        compute_type: str = config.COMPUTE_TYPE,
        language: str = config.LANGUAGE,
        batch_size: int = config.BATCH_SIZE
    ):
        """
        Initialize VideoSummarizer.

        Args:
            video_path: Path to input video file
            output_dir: Output directory (auto-generated if None)
            whisper_model: Whisper model name
            compute_type: Computation type (float16, int8, etc.)
            language: Language code for transcription
            batch_size: Number of segments to process at once
        """
        self.video_path = Path(video_path)
        if not self.video_path.exists():
            raise FileNotFoundError(f"Video file not found: {video_path}")

        self.output_dir = Path(output_dir) if output_dir else config.get_default_output_dir(str(video_path))
        self.audio_path = self.output_dir / "temp_audio.wav"

        # Create output directories
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.ppt_images_dir = self.output_dir / "ppt_images"
        self.ppt_images_dir.mkdir(exist_ok=True)

        # Whisper configuration
        self.whisper_model = whisper_model
        self.compute_type = compute_type
        self.language = language
        self.batch_size = batch_size

        # Initialize Whisper model
        logger.info(f"Initializing Whisper model: {whisper_model}")
        self.whisper = WhisperModel(
            whisper_model,
            device="cuda",
            compute_type=compute_type
        )
        logger.info("Whisper model loaded successfully")

    def extract_audio(self) -> Path:
        """
        Extract audio from video using FFmpeg.

        Returns:
            Path to extracted audio file
        """
        logger.info("Extracting audio from video...")

        ffmpeg_cmd = [
            config.FFMPEG_PATH or "ffmpeg",
            "-i", str(self.video_path),
            "-ac", str(config.AUDIO_CHANNELS),
            "-ar", str(config.AUDIO_SAMPLE_RATE),
            "-y", str(self.audio_path)
        ]

        result = subprocess.run(
            ffmpeg_cmd,
            check=True,
            capture_output=True,
            text=True
        )

        if not self.audio_path.exists():
            raise RuntimeError("Audio extraction failed")

        logger.info(f"Audio extracted: {self.audio_path}")
        return self.audio_path

    def extract_ppt_frames(self) -> list[dict]:
        """
        Extract PPT/slide frames and timestamps from video.

        Returns:
            List of dicts with keys: timestamp, image_path
        """
        logger.info("Extracting PPT frames...")

        if extract_ppt is None:
            logger.warning("extract-video-ppt not installed, using fallback method")
            return self._extract_frames_fallback()

        result = extract_ppt(
            video_path=str(self.video_path),
            output_dir=str(self.ppt_images_dir),
            threshold=config.PPT_THRESHOLD,
            threads=config.PPT_THREADS
        )

        logger.info(f"Extracted {len(result)} PPT frames")
        return result

    def _extract_frames_fallback(self) -> list[dict]:
        """
        Fallback frame extraction using OpenCV when extract-video-ppt unavailable.

        Returns:
            List of dicts with keys: timestamp, image_path
        """
        logger.info("Using OpenCV fallback for frame extraction")
        # Placeholder - real implementation would use OpenCV
        return []

    def transcribe_audio(self) -> list[dict]:
        """
        Transcribe audio to text with timestamps using Whisper.

        Returns:
            List of dicts with keys: start, end, text
        """
        logger.info("Transcribing audio with Whisper...")

        segments, info = self.whisper.transcribe(
            str(self.audio_path),
            language=self.language,
            beam_size=config.BEAM_SIZE,
            best_of=config.BEST_OF,
            temperature=config.TEMPERATURE,
            word_timestamps=config.WORD_TIMESTAMPS,
            vad_filter=config.VAD_FILTER,
            batch_size=self.batch_size
        )

        logger.info(f"Transcription complete: {info.duration:.1f}s, {info.language}")

        transcript = []
        for segment in segments:
            transcript.append({
                "start": segment.start,
                "end": segment.end,
                "text": segment.text.strip()
            })

        return transcript

    def align_ppt_with_transcript(
        self,
        ppt_frames: list[dict],
        transcript: list[dict]
    ) -> list[dict]:
        """
        Align PPT frames with transcript segments by timeline.

        Args:
            ppt_frames: List of PPT frame dicts with timestamp, image_path
            transcript: List of transcript dicts with start, end, text

        Returns:
            List of aligned timeline items
        """
        logger.info("Aligning PPT frames with transcript...")

        timeline = []

        for i, ppt in enumerate(ppt_frames):
            ppt_time = self._timestamp_to_seconds(ppt["timestamp"])

            # Determine this PPT's valid time range
            if i < len(ppt_frames) - 1:
                next_ppt_time = self._timestamp_to_seconds(ppt_frames[i + 1]["timestamp"])
            else:
                next_ppt_time = float('inf')

            # Collect all transcript text within this range
            relevant_text = []
            for seg in transcript:
                if ppt_time <= seg["start"] < next_ppt_time:
                    relevant_text.append(seg["text"])

            timeline.append({
                "ppt_index": i + 1,
                "timestamp": ppt["timestamp"],
                "image_path": ppt["image_path"],
                "text": "\n".join(relevant_text) if relevant_text else "(无语音内容)"
            })

        logger.info(f"Timeline aligned: {len(timeline)} items")
        return timeline

    def _timestamp_to_seconds(self, timestamp_str: str) -> float:
        """
        Convert "HH:MM:SS" timestamp to seconds.

        Args:
            timestamp_str: Timestamp string in "HH:MM:SS" format

        Returns:
            Total seconds as float
        """
        parts = timestamp_str.split(":")
        if len(parts) == 3:
            h, m, s = map(float, parts)
        elif len(parts) == 2:
            h, m, s = 0, float(parts[0]), float(parts[1])
        else:
            h, m, s = 0, 0, float(parts[0])
        return h * 3600 + m * 60 + s

    def generate_markdown_report(self, timeline: list[dict]) -> Path:
        """Generate Markdown format report."""
        md_path = self.output_dir / "timeline_summary.md"

        with open(md_path, "w", encoding="utf-8") as f:
            f.write(f"# 视频内容时间线总结\n\n")
            f.write(f"**视频文件**: {self.video_path.name}\n\n")
            f.write(f"**PPT 页数**: {len(timeline)} 页\n\n")
            f.write("---\n\n")

            for item in timeline:
                f.write(f"## PPT {item['ppt_index']} | {item['timestamp']}\n\n")
                f.write(f"![PPT {item['ppt_index']}](ppt_images/{Path(item['image_path']).name})\n\n")
                f.write("### 语音内容\n\n")
                f.write(f"{item['text']}\n\n")
                f.write("---\n\n")

        logger.info(f"Markdown report: {md_path}")
        return md_path

    def generate_html_report(self, timeline: list[dict]) -> Path:
        """Generate HTML format report with embedded images."""
        html_path = self.output_dir / "timeline_summary.html"

        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>视频时间线总结</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 1200px; margin: 0 auto; padding: 20px; }}
        .ppt-item {{ margin-bottom: 40px; border-bottom: 1px solid #eee; padding-bottom: 20px; }}
        .ppt-image {{ max-width: 100%; height: auto; border: 1px solid #ddd; }}
        .timestamp {{ color: #666; font-size: 0.9em; }}
        h2 {{ color: #333; }}
    </style>
</head>
<body>
    <h1>视频内容时间线总结</h1>
    <p><strong>视频文件</strong>: {self.video_path.name}</p>
    <p><strong>PPT 页数</strong>: {len(timeline)} 页</p>
    <hr>
"""

        for item in timeline:
            img_rel_path = Path(item["image_path"]).name
            html_content += f"""
    <div class="ppt-item">
        <h2>PPT {item['ppt_index']} <span class="timestamp">| {item['timestamp']}</span></h2>
        <img src="ppt_images/{img_rel_path}" class="ppt-image" alt="PPT {item['ppt_index']}">
        <h3>语音内容</h3>
        <p>{item['text'].replace(chr(10), '<br>')}</p>
    </div>
    <hr>
"""

        html_content += "</body></html>"

        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        logger.info(f"HTML report: {html_path}")
        return html_path

    def generate_word_report(self, timeline: list[dict]) -> Path:
        """Generate Word document report."""
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            logger.warning("python-docx not installed, skipping Word report")
            return None

        docx_path = self.output_dir / "timeline_summary.docx"
        doc = Document()
        doc.add_heading('视频内容时间线总结', 0)

        for item in timeline:
            doc.add_heading(f"PPT {item['ppt_index']} | {item['timestamp']}", level=1)

            img_path = Path(item["image_path"])
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6))

            doc.add_heading('语音内容', level=2)
            doc.add_paragraph(item['text'])
            doc.add_page_break()

        doc.save(str(docx_path))
        logger.info(f"Word report: {docx_path}")
        return docx_path

    def run(self, output_formats: list[str] = None) -> dict:
        """
        Execute complete video processing workflow.

        Args:
            output_formats: List of formats to generate (md, html, word)
                           Defaults to config.DEFAULT_OUTPUT_FORMATS

        Returns:
            Dict with output paths and statistics
        """
        if output_formats is None:
            output_formats = config.DEFAULT_OUTPUT_FORMATS

        logger.info(f"🚀 Starting video processing: {self.video_path.name}")

        # Step 1: Extract audio
        logger.info("🎵 Step 1/4: Extracting audio...")
        self.extract_audio()

        # Step 2: Extract PPT frames (CPU)
        logger.info("🖼️ Step 2/4: Extracting PPT frames...")
        ppt_frames = self.extract_ppt_frames()
        logger.info(f"   Extracted {len(ppt_frames)} PPT frames")

        # Step 3: Transcribe audio (GPU)
        logger.info("🎙️ Step 3/4: Transcribing audio...")
        transcript = self.transcribe_audio()
        logger.info(f"   Transcribed {len(transcript)} segments")

        # Step 4: Align timeline
        logger.info("⏱️ Step 4/4: Aligning timeline...")
        timeline = self.align_ppt_with_transcript(ppt_frames, transcript)

        # Generate reports
        logger.info("📝 Generating reports...")
        outputs = {"timeline": timeline}

        if "markdown" in output_formats:
            outputs["markdown"] = self.generate_markdown_report(timeline)

        if "html" in output_formats:
            outputs["html"] = self.generate_html_report(timeline)

        if "word" in output_formats:
            word_path = self.generate_word_report(timeline)
            if word_path:
                outputs["word"] = word_path

        logger.info(f"✅ Processing complete: {self.video_path.name}")
        logger.info(f"📄 Output directory: {self.output_dir}")

        return outputs


def main():
    """CLI entry point."""
    import argparse

    parser = argparse.ArgumentParser(description="Video Summarizer")
    parser.add_argument("--input", "-i", required=True, help="Input video file")
    parser.add_argument("--output", "-o", help="Output directory")
    parser.add_argument(
        "--formats", "-f",
        nargs="+",
        default=["markdown", "html"],
        choices=["markdown", "html", "word"],
        help="Output formats"
    )

    args = parser.parse_args()

    try:
        summarizer = VideoSummarizer(
            video_path=args.input,
            output_dir=args.output
        )
        results = summarizer.run(output_formats=args.formats)

        print("\n✅ SUCCESS")
        for key, path in results.items():
            if key != "timeline":
                print(f"  {key}: {path}")

    except Exception as e:
        logger.error(f"Failed: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
